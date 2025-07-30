"""
Export manager for transcription data with multiple format support.

This module provides the ExportManager class which handles exporting
transcription data in various formats including text, markdown, JSON,
CSV, DOCX, and PDF with comprehensive filtering and metadata options.
"""

import csv
import json
import os
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


class ExportFormat:
    """Enum-like class for export formats."""
    TXT = "txt"
    MD = "md"
    JSON = "json"
    CSV = "csv"
    DOCX = "docx"
    PDF = "pdf"


class ExportOptions:
    """Configuration options for export operations."""
    
    def __init__(self,
                 include_metadata: bool = True,
                 include_timestamps: bool = True,
                 include_raw_text: bool = True,
                 include_formatted_text: bool = True,
                 include_file_paths: bool = False,
                 include_working_dir: bool = False,
                 date_format: str = "%Y-%m-%d %H:%M:%S",
                 encoding: str = "utf-8"):
        """
        Initialize export options.
        
        Args:
            include_metadata (bool): Whether to include metadata in export.
            include_timestamps (bool): Whether to include timestamps.
            include_raw_text (bool): Whether to include raw transcription text.
            include_formatted_text (bool): Whether to include AI-formatted text.
            include_file_paths (bool): Whether to include file paths.
            include_working_dir (bool): Whether to include working directory.
            date_format (str): Format string for dates.
            encoding (str): Text encoding for export files.
        """
        self.include_metadata = include_metadata
        self.include_timestamps = include_timestamps
        self.include_raw_text = include_raw_text
        self.include_formatted_text = include_formatted_text
        self.include_file_paths = include_file_paths
        self.include_working_dir = include_working_dir
        self.date_format = date_format
        self.encoding = encoding


class ExportManager:
    """Manager class for exporting transcription data in multiple formats."""
    
    def __init__(self, transcript_manager=None, history_manager=None):
        """
        Initialize the export manager.
        
        Args:
            transcript_manager: TranscriptManager instance for file operations.
            history_manager: HistoryManager instance for accessing history data.
        """
        self.transcript_manager = transcript_manager
        self.history_manager = history_manager
        
    def get_supported_formats(self) -> List[str]:
        """
        Get list of supported export formats.
        
        Returns:
            List[str]: List of supported format extensions.
        """
        formats = [ExportFormat.TXT, ExportFormat.MD, ExportFormat.JSON, ExportFormat.CSV]
        
        if DOCX_AVAILABLE:
            formats.append(ExportFormat.DOCX)
        if PDF_AVAILABLE:
            formats.append(ExportFormat.PDF)
            
        return formats
        
    def export_single_transcription(self, 
                                   text: str,
                                   formatted_text: Optional[str] = None,
                                   timestamp: Optional[str] = None,
                                   metadata: Optional[Dict[str, Any]] = None,
                                   output_path: Optional[Path] = None,
                                   format_type: str = ExportFormat.TXT,
                                   options: Optional[ExportOptions] = None) -> Path:
        """
        Export a single transcription to the specified format.
        
        Args:
            text (str): Raw transcription text.
            formatted_text (Optional[str]): AI-formatted text if available.
            timestamp (Optional[str]): Timestamp of the transcription.
            metadata (Optional[Dict[str, Any]]): Additional metadata.
            output_path (Optional[Path]): Custom output path.
            format_type (str): Export format type.
            options (Optional[ExportOptions]): Export configuration options.
            
        Returns:
            Path: Path to the exported file.
        """
        if options is None:
            options = ExportOptions()
            
        # Generate default filename if no output path provided
        if output_path is None:
            timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"transcription_{timestamp_str}.{format_type}"
            output_path = Path(filename)
            
        # Ensure output directory exists
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Prepare transcription data
        transcription_data = {
            "text": text,
            "formatted_text": formatted_text,
            "timestamp": timestamp or datetime.now().isoformat(),
            "metadata": metadata or {}
        }
        
        # Export based on format
        if format_type == ExportFormat.TXT:
            return self._export_txt_single(transcription_data, output_path, options)
        elif format_type == ExportFormat.MD:
            return self._export_md_single(transcription_data, output_path, options)
        elif format_type == ExportFormat.JSON:
            return self._export_json_single(transcription_data, output_path, options)
        elif format_type == ExportFormat.CSV:
            return self._export_csv_single(transcription_data, output_path, options)
        elif format_type == ExportFormat.DOCX and DOCX_AVAILABLE:
            return self._export_docx_single(transcription_data, output_path, options)
        elif format_type == ExportFormat.PDF and PDF_AVAILABLE:
            return self._export_pdf_single(transcription_data, output_path, options)
        else:
            raise ValueError(f"Unsupported format: {format_type}")
            
    def export_history(self,
                      history_data: List[Dict[str, Any]],
                      output_path: Optional[Path] = None,
                      format_type: str = ExportFormat.JSON,
                      options: Optional[ExportOptions] = None,
                      filter_criteria: Optional[Dict[str, Any]] = None) -> Path:
        """
        Export history data to the specified format.
        
        Args:
            history_data (List[Dict[str, Any]]): List of history entries.
            output_path (Optional[Path]): Custom output path.
            format_type (str): Export format type.
            options (Optional[ExportOptions]): Export configuration options.
            filter_criteria (Optional[Dict[str, Any]]): Filtering criteria.
            
        Returns:
            Path: Path to the exported file.
        """
        if options is None:
            options = ExportOptions()
            
        # Apply filters if provided
        if filter_criteria:
            history_data = self._apply_filters(history_data, filter_criteria)
            
        # Generate default filename if no output path provided
        if output_path is None:
            timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"transcription_history_{timestamp_str}.{format_type}"
            output_path = Path(filename)
            
        # Ensure output directory exists
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Export based on format
        if format_type == ExportFormat.TXT:
            return self._export_txt_history(history_data, output_path, options)
        elif format_type == ExportFormat.MD:
            return self._export_md_history(history_data, output_path, options)
        elif format_type == ExportFormat.JSON:
            return self._export_json_history(history_data, output_path, options)
        elif format_type == ExportFormat.CSV:
            return self._export_csv_history(history_data, output_path, options)
        elif format_type == ExportFormat.DOCX and DOCX_AVAILABLE:
            return self._export_docx_history(history_data, output_path, options)
        elif format_type == ExportFormat.PDF and PDF_AVAILABLE:
            return self._export_pdf_history(history_data, output_path, options)
        else:
            raise ValueError(f"Unsupported format: {format_type}")
            
    def export_session_data(self,
                           session_entries: List[Dict[str, Any]],
                           output_path: Optional[Path] = None,
                           format_type: str = ExportFormat.JSON,
                           options: Optional[ExportOptions] = None) -> Path:
        """
        Export current session transcription data.
        
        Args:
            session_entries (List[Dict[str, Any]]): List of session entries.
            output_path (Optional[Path]): Custom output path.
            format_type (str): Export format type.
            options (Optional[ExportOptions]): Export configuration options.
            
        Returns:
            Path: Path to the exported file.
        """
        if options is None:
            options = ExportOptions()
            
        # Generate default filename if no output path provided
        if output_path is None:
            timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"session_transcriptions_{timestamp_str}.{format_type}"
            output_path = Path(filename)
            
        # Ensure output directory exists
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Convert session entries to standard format
        formatted_entries = []
        for entry in session_entries:
            formatted_entry = {
                "text": getattr(entry, 'text', ''),
                "formatted_text": getattr(entry, 'formatted_text', None),
                "timestamp": getattr(entry, 'timestamp', ''),
                "metadata": {
                    "entry_id": getattr(entry, 'entry_id', ''),
                    "session_export": True
                }
            }
            formatted_entries.append(formatted_entry)
            
        # Export based on format
        if format_type == ExportFormat.TXT:
            return self._export_txt_history(formatted_entries, output_path, options)
        elif format_type == ExportFormat.MD:
            return self._export_md_history(formatted_entries, output_path, options)
        elif format_type == ExportFormat.JSON:
            return self._export_json_history(formatted_entries, output_path, options)
        elif format_type == ExportFormat.CSV:
            return self._export_csv_history(formatted_entries, output_path, options)
        elif format_type == ExportFormat.DOCX and DOCX_AVAILABLE:
            return self._export_docx_history(formatted_entries, output_path, options)
        elif format_type == ExportFormat.PDF and PDF_AVAILABLE:
            return self._export_pdf_history(formatted_entries, output_path, options)
        else:
            raise ValueError(f"Unsupported format: {format_type}")
            
    def _apply_filters(self, data: List[Dict[str, Any]], filters: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Apply filtering criteria to data."""
        filtered_data = data
        
        # Date range filter
        if 'date_from' in filters or 'date_to' in filters:
            date_from = filters.get('date_from')
            date_to = filters.get('date_to')
            
            def date_filter(entry):
                timestamp = entry.get('timestamp', '')
                if isinstance(timestamp, str):
                    try:
                        entry_date = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
                        if date_from and entry_date < date_from:
                            return False
                        if date_to and entry_date > date_to:
                            return False
                    except ValueError:
                        pass
                return True
                
            filtered_data = [entry for entry in filtered_data if date_filter(entry)]
            
        # Working directory filter
        if 'working_dir' in filters:
            working_dir = filters['working_dir']
            filtered_data = [entry for entry in filtered_data 
                           if entry.get('working_dir') == working_dir]
            
        # Text search filter
        if 'text_search' in filters:
            search_term = filters['text_search'].lower()
            filtered_data = [entry for entry in filtered_data 
                           if search_term in entry.get('text', '').lower()]
            
        return filtered_data
        
    def _export_txt_single(self, data: Dict[str, Any], output_path: Path, options: ExportOptions) -> Path:
        """Export single transcription as TXT."""
        content = []
        
        if options.include_metadata and options.include_timestamps:
            content.append(f"Timestamp: {data['timestamp']}")
            content.append("-" * 50)
            
        if options.include_raw_text and data.get('text'):
            content.append("Raw Transcription:")
            content.append(data['text'])
            content.append("")
            
        if options.include_formatted_text and data.get('formatted_text'):
            content.append("AI-Formatted Transcription:")
            content.append(data['formatted_text'])
            content.append("")
            
        with open(output_path, 'w', encoding=options.encoding) as f:
            f.write('\n'.join(content))
            
        return output_path
        
    def _export_md_single(self, data: Dict[str, Any], output_path: Path, options: ExportOptions) -> Path:
        """Export single transcription as Markdown."""
        content = []
        
        if options.include_metadata and options.include_timestamps:
            content.append(f"# Transcription - {data['timestamp']}")
            content.append("")
            
        if options.include_raw_text and data.get('text'):
            content.append("## Raw Transcription")
            content.append(data['text'])
            content.append("")
            
        if options.include_formatted_text and data.get('formatted_text'):
            content.append("## AI-Formatted Transcription")
            content.append(data['formatted_text'])
            content.append("")
            
        with open(output_path, 'w', encoding=options.encoding) as f:
            f.write('\n'.join(content))
            
        return output_path
        
    def _export_json_single(self, data: Dict[str, Any], output_path: Path, options: ExportOptions) -> Path:
        """Export single transcription as JSON."""
        export_data = {}
        
        if options.include_timestamps:
            export_data['timestamp'] = data['timestamp']
            
        if options.include_raw_text and data.get('text'):
            export_data['raw_text'] = data['text']
            
        if options.include_formatted_text and data.get('formatted_text'):
            export_data['formatted_text'] = data['formatted_text']
            
        if options.include_metadata and data.get('metadata'):
            export_data['metadata'] = data['metadata']
            
        with open(output_path, 'w', encoding=options.encoding) as f:
            json.dump(export_data, f, indent=2, ensure_ascii=False)
            
        return output_path
        
    def _export_csv_single(self, data: Dict[str, Any], output_path: Path, options: ExportOptions) -> Path:
        """Export single transcription as CSV."""
        fieldnames = []
        row_data = {}
        
        if options.include_timestamps:
            fieldnames.append('timestamp')
            row_data['timestamp'] = data['timestamp']
            
        if options.include_raw_text and data.get('text'):
            fieldnames.append('raw_text')
            row_data['raw_text'] = data['text']
            
        if options.include_formatted_text and data.get('formatted_text'):
            fieldnames.append('formatted_text')
            row_data['formatted_text'] = data['formatted_text']
            
        with open(output_path, 'w', newline='', encoding=options.encoding) as f:
            writer = csv.DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerow(row_data)
            
        return output_path
        
    def _export_docx_single(self, data: Dict[str, Any], output_path: Path, options: ExportOptions) -> Path:
        """Export single transcription as DOCX."""
        doc = Document()
        
        if options.include_metadata and options.include_timestamps:
            doc.add_heading(f"Transcription - {data['timestamp']}", 0)
            
        if options.include_raw_text and data.get('text'):
            doc.add_heading("Raw Transcription", level=1)
            doc.add_paragraph(data['text'])
            
        if options.include_formatted_text and data.get('formatted_text'):
            doc.add_heading("AI-Formatted Transcription", level=1)
            doc.add_paragraph(data['formatted_text'])
            
        doc.save(output_path)
        return output_path
        
    def _export_pdf_single(self, data: Dict[str, Any], output_path: Path, options: ExportOptions) -> Path:
        """Export single transcription as PDF."""
        doc = SimpleDocTemplate(str(output_path), pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        if options.include_metadata and options.include_timestamps:
            title = Paragraph(f"Transcription - {data['timestamp']}", styles['Title'])
            story.append(title)
            story.append(Spacer(1, 12))
            
        if options.include_raw_text and data.get('text'):
            heading = Paragraph("Raw Transcription", styles['Heading1'])
            story.append(heading)
            content = Paragraph(data['text'], styles['Normal'])
            story.append(content)
            story.append(Spacer(1, 12))
            
        if options.include_formatted_text and data.get('formatted_text'):
            heading = Paragraph("AI-Formatted Transcription", styles['Heading1'])
            story.append(heading)
            content = Paragraph(data['formatted_text'], styles['Normal'])
            story.append(content)
            
        doc.build(story)
        return output_path
        
    def _export_txt_history(self, history_data: List[Dict[str, Any]], output_path: Path, options: ExportOptions) -> Path:
        """Export history as TXT."""
        content = []
        content.append("Transcription History Export")
        content.append("=" * 50)
        content.append("")
        
        for i, entry in enumerate(history_data, 1):
            content.append(f"Entry {i}")
            content.append("-" * 20)
            
            if options.include_timestamps:
                timestamp = entry.get('timestamp', 'Unknown')
                content.append(f"Timestamp: {timestamp}")
                
            if options.include_raw_text and entry.get('text'):
                content.append("Text:")
                content.append(entry['text'])
                content.append("")
                
            if options.include_formatted_text and entry.get('formatted_text'):
                content.append("Formatted Text:")
                content.append(entry['formatted_text'])
                content.append("")
                
            content.append("")
            
        with open(output_path, 'w', encoding=options.encoding) as f:
            f.write('\n'.join(content))
            
        return output_path
        
    def _export_md_history(self, history_data: List[Dict[str, Any]], output_path: Path, options: ExportOptions) -> Path:
        """Export history as Markdown."""
        content = []
        content.append("# Transcription History Export")
        content.append("")
        
        for i, entry in enumerate(history_data, 1):
            content.append(f"## Entry {i}")
            
            if options.include_timestamps:
                timestamp = entry.get('timestamp', 'Unknown')
                content.append(f"**Timestamp:** {timestamp}")
                content.append("")
                
            if options.include_raw_text and entry.get('text'):
                content.append("### Raw Transcription")
                content.append(entry['text'])
                content.append("")
                
            if options.include_formatted_text and entry.get('formatted_text'):
                content.append("### AI-Formatted Transcription")
                content.append(entry['formatted_text'])
                content.append("")
                
            content.append("---")
            content.append("")
            
        with open(output_path, 'w', encoding=options.encoding) as f:
            f.write('\n'.join(content))
            
        return output_path
        
    def _export_json_history(self, history_data: List[Dict[str, Any]], output_path: Path, options: ExportOptions) -> Path:
        """Export history as JSON."""
        export_data = {
            "export_timestamp": datetime.now().isoformat(),
            "entries": []
        }
        
        for entry in history_data:
            export_entry = {}
            
            if options.include_timestamps:
                export_entry['timestamp'] = entry.get('timestamp', '')
                
            if options.include_raw_text and entry.get('text'):
                export_entry['text'] = entry['text']
                
            if options.include_formatted_text and entry.get('formatted_text'):
                export_entry['formatted_text'] = entry['formatted_text']
                
            if options.include_metadata and entry.get('metadata'):
                export_entry['metadata'] = entry['metadata']
                
            export_data['entries'].append(export_entry)
            
        with open(output_path, 'w', encoding=options.encoding) as f:
            json.dump(export_data, f, indent=2, ensure_ascii=False)
            
        return output_path
        
    def _export_csv_history(self, history_data: List[Dict[str, Any]], output_path: Path, options: ExportOptions) -> Path:
        """Export history as CSV."""
        fieldnames = []
        
        if options.include_timestamps:
            fieldnames.append('timestamp')
        if options.include_raw_text:
            fieldnames.append('raw_text')
        if options.include_formatted_text:
            fieldnames.append('formatted_text')
            
        with open(output_path, 'w', newline='', encoding=options.encoding) as f:
            writer = csv.DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()
            
            for entry in history_data:
                row_data = {}
                
                if options.include_timestamps:
                    row_data['timestamp'] = entry.get('timestamp', '')
                if options.include_raw_text:
                    row_data['raw_text'] = entry.get('text', '')
                if options.include_formatted_text:
                    row_data['formatted_text'] = entry.get('formatted_text', '')
                    
                writer.writerow(row_data)
                
        return output_path
        
    def _export_docx_history(self, history_data: List[Dict[str, Any]], output_path: Path, options: ExportOptions) -> Path:
        """Export history as DOCX."""
        doc = Document()
        doc.add_heading("Transcription History Export", 0)
        
        for i, entry in enumerate(history_data, 1):
            doc.add_heading(f"Entry {i}", level=1)
            
            if options.include_timestamps:
                timestamp = entry.get('timestamp', 'Unknown')
                doc.add_paragraph(f"Timestamp: {timestamp}")
                
            if options.include_raw_text and entry.get('text'):
                doc.add_heading("Raw Transcription", level=2)
                doc.add_paragraph(entry['text'])
                
            if options.include_formatted_text and entry.get('formatted_text'):
                doc.add_heading("AI-Formatted Transcription", level=2)
                doc.add_paragraph(entry['formatted_text'])
                
            doc.add_page_break()
            
        doc.save(output_path)
        return output_path
        
    def _export_pdf_history(self, history_data: List[Dict[str, Any]], output_path: Path, options: ExportOptions) -> Path:
        """Export history as PDF."""
        doc = SimpleDocTemplate(str(output_path), pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        title = Paragraph("Transcription History Export", styles['Title'])
        story.append(title)
        story.append(Spacer(1, 12))
        
        for i, entry in enumerate(history_data, 1):
            heading = Paragraph(f"Entry {i}", styles['Heading1'])
            story.append(heading)
            
            if options.include_timestamps:
                timestamp = entry.get('timestamp', 'Unknown')
                timestamp_p = Paragraph(f"Timestamp: {timestamp}", styles['Normal'])
                story.append(timestamp_p)
                story.append(Spacer(1, 6))
                
            if options.include_raw_text and entry.get('text'):
                raw_heading = Paragraph("Raw Transcription", styles['Heading2'])
                story.append(raw_heading)
                raw_content = Paragraph(entry['text'], styles['Normal'])
                story.append(raw_content)
                story.append(Spacer(1, 6))
                
            if options.include_formatted_text and entry.get('formatted_text'):
                formatted_heading = Paragraph("AI-Formatted Transcription", styles['Heading2'])
                story.append(formatted_heading)
                formatted_content = Paragraph(entry['formatted_text'], styles['Normal'])
                story.append(formatted_content)
                
            story.append(Spacer(1, 12))
            
        doc.build(story)
        return output_path